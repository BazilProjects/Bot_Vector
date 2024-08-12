import warnings
warnings.filterwarnings("ignore")
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split, learning_curve
from sklearn.linear_model import LinearRegression
from sklearn.tree import DecisionTreeRegressor
from sklearn.ensemble import RandomForestRegressor, ExtraTreesRegressor
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import PolynomialFeatures, OneHotEncoder
from sklearn.compose import ColumnTransformer
import plotly.graph_objs as go
import plotly.io as pio
from itertools import combinations
from metaapi_cloud_sdk import MetaApi
import asyncio
import joblib
import sklearn
import os
from docx import Document
from sklearn.model_selection import cross_val_score


from datetime import datetime, timedelta
token = os.getenv('TOKEN') or 'eyJhbGciOiJSUzUxMiIsInR5cCI6IkpXVCJ9.eyJfaWQiOiJlNWZkNGRkYWZmZmIyMDk2YTAyMWYzNjZiY2YxYjYwYSIsInBlcm1pc3Npb25zIjpbXSwiYWNjZXNzUnVsZXMiOlt7ImlkIjoidHJhZGluZy1hY2NvdW50LW1hbmFnZW1lbnQtYXBpIiwibWV0aG9kcyI6WyJ0cmFkaW5nLWFjY291bnQtbWFuYWdlbWVudC1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZXN0LWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1ycGMtYXBpIiwibWV0aG9kcyI6WyJtZXRhYXBpLWFwaTp3czpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZWFsLXRpbWUtc3RyZWFtaW5nLWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6d3M6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6Im1ldGFzdGF0cy1hcGkiLCJtZXRob2RzIjpbIm1ldGFzdGF0cy1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoicmlzay1tYW5hZ2VtZW50LWFwaSIsIm1ldGhvZHMiOlsicmlzay1tYW5hZ2VtZW50LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJjb3B5ZmFjdG9yeS1hcGkiLCJtZXRob2RzIjpbImNvcHlmYWN0b3J5LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJtdC1tYW5hZ2VyLWFwaSIsIm1ldGhvZHMiOlsibXQtbWFuYWdlci1hcGk6cmVzdDpkZWFsaW5nOio6KiIsIm10LW1hbmFnZXItYXBpOnJlc3Q6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6ImJpbGxpbmctYXBpIiwibWV0aG9kcyI6WyJiaWxsaW5nLWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19XSwidG9rZW5JZCI6IjIwMjEwMjEzIiwiaW1wZXJzb25hdGVkIjpmYWxzZSwicmVhbFVzZXJJZCI6ImU1ZmQ0ZGRhZmZmYjIwOTZhMDIxZjM2NmJjZjFiNjBhIiwiaWF0IjoxNzE4NjAwMTc5fQ.P_Q7S9gllY-A0ygrF72pmpyUrno0VZ6_gBXIT31fLOwKFcEEeqopfcEH7yUL6upwTe69YAitfuy3OcDfjBNL7D7Vnuh1FUUUb2EbuGEnBi-B3GSazhZ83uSeAa89zutuNsr7DrptYf-ZHIUw10NSbIHZjhTKOsT9GoV-lv6QLsaxe87wJ8hbv5ajWvo2VKYhJKv0fQFDP2SwsboGnjC4ioqMGoGFhAv2BNLs3nXZL8SWe0tEYIcLmfWB1sWTlPDrsnaOzRdyOzUwbRljoJzns1BUyHW375eQTR93oNjN_P1zJNL9J_V8rqc-nIJcQPhIFytwaFDy-Z3DJHZQJk6mAbDLkphjQAvGeXrwjjk_uQNuY0WzmJmv21dFeQ3aFJhW-wH7l_KYEYjwqxP_H2lU0_AXpiGxn0ZWOoRbyp2Uer8X2hBg_psUd7RlhmfVqXQqTkJLpFpZrP980rq1S_deFWWiZPihXqkzTsFwUTl2DWMyuTxwsQnCBi-Dbt11XY-1kXe7uABqL9L6YbaclohAUVQF1pIVONbELTkTzUqfzqf30TLsxT67xScvd51-smHTDBYyvSDjzWLQ-HCl9gMJlZ43NcJtA9fckE47yf7O5YaaaA3QXhNSO86jiDvJ_spVijZXi62P_d4d71fCpUZTum1WSwRAiF5scvIjwFMSL98'
accountId = os.getenv('ACCOUNT_ID') or '653d65c4-a70f-49ac-a6de-deea63238808'

symbol_list = [
    'XAUUSDm',  # Gold/US Dollar (Commodity)
    'GBPAUDm' ,  # British Pound/Australian Dollar (Minor)
    'BTCUSDm',
    'EURUSDm',  # Euro/US Dollar (Major)
    'GBPUSDm',  # British Pound/US Dollar (Major)
    

]

data = {}
timeframe='15m'
pages=5
n_estimators=1
min_samples_leaf=1
shuffle=True
max_depth=50
test_size=0.01

def decimal_places(number):
    # Convert the number to a string
    num_str = str(number)
    
    # Check if there is a decimal point
    if '.' in num_str:
        # Find the index of the decimal point
        decimal_index = num_str.index('.')
        
        # Count the characters after the decimal point
        num_decimal_places = len(num_str) - decimal_index - 1
        
        return num_decimal_places
    else:
        # If there is no decimal point, return 0
        return 0


# Function to add nested dictionary content to a .docx file
def add_nested_dict_to_docx(info_dict, filename=f'output{pages}{timeframe}.docx'):
    doc = Document()
    doc.add_heading('AI analysis and market forecast For timeframe', 0)
    for symbol, details in info_dict.items():
        doc.add_heading(symbol, level=1)
        for key, value in details.items():
            doc.add_paragraph(f"{key}: {value}", style='List Number')
    doc.save(filename)
    print(f"Information added to {filename} successfully.")

# Function to add data to the dictionary
def add_symbol_data(symbol, r2_2, mse_2, r2_2_1, mse_2_1, r2_2_2, mse_2_2, next_close, next_low, next_high, actual_close, actual_low, actual_high, diff2, diff1):
    if symbol not in data:
        data[symbol] = {}
    
    data[symbol] = {
        'r2_2': r2_2,
        'mse_2': mse_2,
        'r2_2_1': r2_2_1,
        'mse_2_1': mse_2_1,
        'r2_2_2': r2_2_2,
        'mse_2_2': mse_2_2,
        'next_close': next_close,
        'next_low': next_low,
        'next_high': next_high,
        'actual': {
            'close': actual_close,
            'low': actual_low,
            'high': actual_high
        },
        'diff2': diff2,
        'diff1': diff1
    }

def prepare(df):
    #df=add_stop_losse(df)
    df=df.drop(columns=['symbol','timeframe','brokerTime'])

    df['time'] = pd.to_datetime(df['time'])
    df['time'] = df['time'].astype(int)// 10**99 
    # Generate new features and handle data preparation
    df_new =df

    return df_new
async def main2(timeframe,pages):
    print('Up and runing')
    for symbol in symbol_list:
        try:

            
            df=pd.read_csv(f'../COLLECT CANDLES/{timeframe}/{symbol}{timeframe}{str(pages)}.csv')
            actual=df.tail(1)
            df=df.head(len(df)-1)
            

            #print(df)
            if not df.empty:
                print(f'Symbol: {symbol}')
                df_new=prepare(df)
                dt=pd.DataFrame(df_new)
                # Print all combinations
                """
                df_new['Candle_close'] = df_new['close'].shift(-1)
                df_new=df_new.dropna()

                X =df_new.drop(columns=['Candle_close'])
                y= df_new['Candle_close']
                """
                df_new['Candle_close'] = df_new['close'].shift(-1)
                df_new['Canlde_low']=df_new['low'].shift(-1)
                df_new['Canlde_high']=df_new['high'].shift(-1)

                #print(df_new[['Prediction','trade_max','stop_losses','stop_losses_predictions']])
                df_new=df_new.dropna()
                #print(df_new[['Prediction','trade_max','stop_losses','stop_losses_predictions']])
                
                X =df_new.drop(columns=['Candle_close','Canlde_low','Canlde_high'])
                y= df_new[['Candle_close','Canlde_low','Canlde_high']]
                

                
                # Initialize RandomForestRegressor
                model = DecisionTreeRegressor()
                model.fit(X,y)
                joblib.dump((model, sklearn.__version__), f'Regressors/model{symbol}{timeframe}close.pkl')
                # Perform cross-validation with R² as the scoring metric
                scores = cross_val_score(model, X, y, cv=100, scoring='r2')

                # Compute mean and standard deviation of the scores
                mean_score = scores.mean()
                std_score = scores.std()

                print(f"Cross-Validation R²: {mean_score} ± {std_score}")
                print(f' '*500)

        except Exception as e:
            #print(f'{symbol} failed')
            raise e
            pass

asyncio.run(main2(timeframe,pages))