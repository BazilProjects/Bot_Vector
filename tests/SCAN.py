import warnings
warnings.filterwarnings("ignore")
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split, learning_curve
from sklearn.linear_model import LinearRegression
from sklearn.tree import DecisionTreeRegressor
from sklearn.ensemble import RandomForestRegressor, ExtraTreesRegressor,HistGradientBoostingRegressor
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import PolynomialFeatures, OneHotEncoder
from sklearn.compose import ColumnTransformer
import plotly.graph_objs as go
import plotly.io as pio
from itertools import combinations
from metaapi_cloud_sdk import MetaApi
import asyncio
import joblib
import sklearn
import pandas_ta as ta
import os
from docx import Document
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error, max_error

from sklearn.linear_model import Ridge, Lasso, ElasticNet, BayesianRidge, HuberRegressor
from sklearn.svm import SVR
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor, AdaBoostRegressor
from sklearn.neighbors import KNeighborsRegressor
from sklearn.linear_model import LinearRegression
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler

from datetime import datetime, timedelta
token = os.getenv('TOKEN') or 'eyJhbGciOiJSUzUxMiIsInR5cCI6IkpXVCJ9.eyJfaWQiOiJlNWZkNGRkYWZmZmIyMDk2YTAyMWYzNjZiY2YxYjYwYSIsInBlcm1pc3Npb25zIjpbXSwiYWNjZXNzUnVsZXMiOlt7ImlkIjoidHJhZGluZy1hY2NvdW50LW1hbmFnZW1lbnQtYXBpIiwibWV0aG9kcyI6WyJ0cmFkaW5nLWFjY291bnQtbWFuYWdlbWVudC1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZXN0LWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1ycGMtYXBpIiwibWV0aG9kcyI6WyJtZXRhYXBpLWFwaTp3czpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZWFsLXRpbWUtc3RyZWFtaW5nLWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6d3M6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6Im1ldGFzdGF0cy1hcGkiLCJtZXRob2RzIjpbIm1ldGFzdGF0cy1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoicmlzay1tYW5hZ2VtZW50LWFwaSIsIm1ldGhvZHMiOlsicmlzay1tYW5hZ2VtZW50LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJjb3B5ZmFjdG9yeS1hcGkiLCJtZXRob2RzIjpbImNvcHlmYWN0b3J5LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJtdC1tYW5hZ2VyLWFwaSIsIm1ldGhvZHMiOlsibXQtbWFuYWdlci1hcGk6cmVzdDpkZWFsaW5nOio6KiIsIm10LW1hbmFnZXItYXBpOnJlc3Q6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6ImJpbGxpbmctYXBpIiwibWV0aG9kcyI6WyJiaWxsaW5nLWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19XSwidG9rZW5JZCI6IjIwMjEwMjEzIiwiaW1wZXJzb25hdGVkIjpmYWxzZSwicmVhbFVzZXJJZCI6ImU1ZmQ0ZGRhZmZmYjIwOTZhMDIxZjM2NmJjZjFiNjBhIiwiaWF0IjoxNzE4NjAwMTc5fQ.P_Q7S9gllY-A0ygrF72pmpyUrno0VZ6_gBXIT31fLOwKFcEEeqopfcEH7yUL6upwTe69YAitfuy3OcDfjBNL7D7Vnuh1FUUUb2EbuGEnBi-B3GSazhZ83uSeAa89zutuNsr7DrptYf-ZHIUw10NSbIHZjhTKOsT9GoV-lv6QLsaxe87wJ8hbv5ajWvo2VKYhJKv0fQFDP2SwsboGnjC4ioqMGoGFhAv2BNLs3nXZL8SWe0tEYIcLmfWB1sWTlPDrsnaOzRdyOzUwbRljoJzns1BUyHW375eQTR93oNjN_P1zJNL9J_V8rqc-nIJcQPhIFytwaFDy-Z3DJHZQJk6mAbDLkphjQAvGeXrwjjk_uQNuY0WzmJmv21dFeQ3aFJhW-wH7l_KYEYjwqxP_H2lU0_AXpiGxn0ZWOoRbyp2Uer8X2hBg_psUd7RlhmfVqXQqTkJLpFpZrP980rq1S_deFWWiZPihXqkzTsFwUTl2DWMyuTxwsQnCBi-Dbt11XY-1kXe7uABqL9L6YbaclohAUVQF1pIVONbELTkTzUqfzqf30TLsxT67xScvd51-smHTDBYyvSDjzWLQ-HCl9gMJlZ43NcJtA9fckE47yf7O5YaaaA3QXhNSO86jiDvJ_spVijZXi62P_d4d71fCpUZTum1WSwRAiF5scvIjwFMSL98'
accountId = os.getenv('ACCOUNT_ID') or '653d65c4-a70f-49ac-a6de-deea63238808'

symbol_list = [

    'BTCUSDm',

]
"""
    

"""
data = {}
timeframe='5m'
pages=10
n_estimators=1
min_samples_leaf=1
shuffle=True
max_depth=50
test_size=0.01
import statistics
from collections import Counter

import statistics
from collections import Counter

def analyze_list(float_list):
    if not float_list:
        return "List is empty"

    analysis = {}
    
    # Basic Statistics
    analysis['mean'] = sum(float_list) / len(float_list)
    analysis['median'] = statistics.median(float_list)
    analysis['mode'] = statistics.multimode(float_list)
    analysis['std_dev'] = statistics.stdev(float_list)
    analysis['range'] = max(float_list) - min(float_list)
    
    # Count of values based on a custom condition (e.g., greater than a certain threshold)
    analysis['greater_than_50'] = len([x for x in float_list if x > 50.0])
    analysis['less_than_50'] = len([x for x in float_list if x < 50.0])
    
    # Sum of fractional parts
    analysis['sum_of_fractional_parts'] = sum([x % 1 for x in float_list])
    
    return analysis


def is_prime(n):
    if n <= 1:
        return False
    for i in range(2, int(n ** 0.5) + 1):
        if n % i == 0:
            return False
    return True
def decimal_places(number):
    # Convert the number to a string
    num_str = str(number)
    
    # Check if there is a decimal point
    if '.' in num_str:
        # Find the index of the decimal point
        decimal_index = num_str.index('.')
        
        # Count the characters after the decimal point
        num_decimal_places = len(num_str) - decimal_index - 1
        
        return num_decimal_places
    else:
        # If there is no decimal point, return 0
        return 0


# Function to add nested dictionary content to a .docx file
def add_nested_dict_to_docx(info_dict, filename=f'output{pages}{timeframe}.docx'):
    doc = Document()
    doc.add_heading('AI analysis and market forecast For timeframe', 0)
    for symbol, details in info_dict.items():
        doc.add_heading(symbol, level=1)
        for key, value in details.items():
            doc.add_paragraph(f"{key}: {value}", style='List Number')
    doc.save(filename)
    print(f"Information added to {filename} successfully.")

# Function to add data to the dictionary
def add_symbol_data(symbol, r2_2, mse_2, r2_2_1, mse_2_1, r2_2_2, mse_2_2, next_close, next_low, next_high, actual_close, actual_low, actual_high, diff2, diff1):
    if symbol not in data:
        data[symbol] = {}
    
    data[symbol] = {
        'r2_2': r2_2,
        'mse_2': mse_2,
        'r2_2_1': r2_2_1,
        'mse_2_1': mse_2_1,
        'r2_2_2': r2_2_2,
        'mse_2_2': mse_2_2,
        'next_close': next_close,
        'next_low': next_low,
        'next_high': next_high,
        'actual': {
            'close': actual_close,
            'low': actual_low,
            'high': actual_high
        },
        'diff2': diff2,
        'diff1': diff1
    }

def prepare(df):
    #df=add_stop_losse(df)
    df=df.drop(columns=['symbol','timeframe','brokerTime'])

    df['time'] = pd.to_datetime(df['time'])
    df['time'] = df['time'].astype(int)// 10**99
    
    # Generate new features and handle data preparation
    df_new =df

    return df_new
async def main2(timeframe,pages):
    print('Up and runing')
    for symbol in symbol_list:
        print(f'Symbol: {symbol}')
        best_profit_to_loss_ratio=1
        best_rsi_period=1
        best_sma_10=1
        best_sma_30=1
        overall_total_number_of_trades=1
        best_added_lag=0
        if symbol=='XAUUSDm':
            added_lag_lower=1
            added_lag_upper=3
            jump_parameters=1
        elif symbol=='BTCUSDm':
            added_lag_lower=100
            added_lag_upper=200
            jump_parameters=100
        elif symbol=='GBPAUDm':
            added_lag_lower=0.0004
            added_lag_upper=0.00017
            jump_parameters=0.0001
        elif symbol=='GBPUSDm':
            jump_parameters=0.0001
            added_lag_lower=0.0004
            added_lag_upper=0.0015
        elif symbol=='EURUSDm':
            jump_parameters=0.0001
            added_lag_lower=0.0004
            added_lag_upper=0.0009 
        for rsi_period in np.arange(20,35,2):
            for sma_10_period in np.arange(20,35,2):
                for sma_30_period in np.arange(20,50,5):
                    for added_lag in np.arange(added_lag_lower,added_lag_upper,jump_parameters):
        
                        try:               
                            df=pd.read_csv(f'../COLLECT CANDLES/{timeframe}/{symbol}{timeframe}{str(pages)}.csv')

                            df=df.head(5500).reset_index()
                            
                            #print(df)
                            if not df.empty:
                                
                                #df['close'] = df['close'].rolling(window=5).mean()

                                #df['close'] = df['close'].ewm(span=5).mean()
                                
                                df_new=prepare(df)
                                dt=pd.DataFrame(df_new)
                                # Print all combinations
                                """
                                df_new['Candle_close'] = df_new['close'].shift(-1)
                                df_new=df_new.dropna()

                                X =df_new.drop(columns=['Candle_close'])
                                y= df_new['Candle_close']
                                """
                                df_new['Candle_close'] = df_new['close'].shift(-1)
                                df_new['Canlde_low']=df_new['low'].shift(-1)
                                df_new['Canlde_high']=df_new['high'].shift(-1)
                                df_new['Canlde_open']=df_new['open'].shift(-1)
                                #df_new['Failure']=(df_new['close'].std()>29.280487294763798).astype(int)
                                
                                # Create labels: 1 for 'buy' (next close is higher), 0 for 'sell' (next close is lower)
                                df_new['Label_close'] = (df_new['Candle_close'] > df_new['close']).astype(int)
                                df_new['Label_open'] = (df_new['Canlde_open'] > df_new['Candle_close']).astype(int)
                                df_new['RSI'] = ta.rsi(df_new['close'], length=rsi_period)
                                df_new['sma_10'] = ta.sma(df['close'], length=sma_10_period)
                                df_new['sma_30'] = ta.sma(df['close'], length=sma_30_period)
                                # Buy Signal: When sma_10 crosses above sma_30
                                df_new['buy_signal'] = ((df_new['sma_10'] > df_new['sma_30']) & 
                                                        (df_new['sma_10'].shift(1) <= df_new['sma_30'].shift(1))).astype(int)

                                # Sell Signal: When sma_10 crosses below sma_30
                                df_new['sell_signal'] = ((df_new['sma_10'] < df_new['sma_30']) & 
                                                         (df_new['sma_10'].shift(1) >= df_new['sma_30'].shift(1))).astype(int)
                                buy_candles_condition = df_new['close'] > df_new['open']

                                # Calculate stop loss for buy candles only
                                stop_loss_buy = df_new['open'] - df_new['low'] + added_lag
                                stop_loss_buy = stop_loss_buy.where(buy_candles_condition, other=0)

                                sell_candles_condition = df_new['close'] < df_new['open']

                                # Calculate stop loss for sell position (next_high - open - 0.0004)
                                stop_loss_sell = df_new['high'] - df_new['open'] + added_lag
                                stop_loss_sell = stop_loss_sell.where(sell_candles_condition, other=0)
                                df_new['stop_losses']=stop_loss_buy+stop_loss_sell
                                df_new['stop_losses']==df_new['stop_losses'].shift(-1)
                                #print(df_new[['Prediction','trade_max','stop_losses','stop_losses_predictions']])
                                df_new=df_new.dropna()
                                

                                X =df_new.drop(columns=['Candle_close','Canlde_low','Canlde_high','Canlde_open','Label_open','Label_close','stop_losses'])
                                y= df_new[['Candle_close','stop_losses']]#'Canlde_low','Canlde_high','Label_close','Label_open']]#
                                
                                from sklearn.preprocessing import MinMaxScaler

                                # Initialize the MinMaxScaler
                                scaler = MinMaxScaler()
                                X = scaler.fit_transform(X)
                                #y=pd.DataFrame(scaler.fit_transform(y), columns=y.columns)
                                # Split the data into training and testing sets
                                X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42, shuffle=False)
                                over_aspan_of=len(y_test)
                                
                                # Initialize RandomForestRegressor
                                model = DecisionTreeRegressor()#RandomForestRegressor(n_estimators=1, min_samples_leaf=2, max_depth=20, random_state=42)
                                model.fit(X_train, y_train)
                            
                                # Predict on the test set
                                y_pred = model.predict(X_test)

                                # Calculate metrics

                                total_expected_trade=1
                                total_expected_wins=1
                                total_expected_losses=1
                                biggest_buy_stop_loss=1
                                biggest_sell_stop_loss=1
                                for i in range(2,y_pred.shape[0]):
                                    pred_close=round(float(y_pred[i][0]),decimal_places(df.iloc[i]['close']))
                                    current_close =y_test.iloc[i][0]
                                    previous_close=y_test.iloc[i-1][0]# df_new.iloc[i-1]['close']
                                    current_open= df_new.iloc[i]['open']
                                    current_low= df_new.iloc[i]['low']
                                    current_high= df_new.iloc[i]['high']
                                    rsi=df_new.iloc[i-1]['RSI']
                                    sma_10=df_new.iloc[i-1]['sma_10']
                                    sma_30=df_new.iloc[i-1]['sma_30']
                                    stop_loss_away=float(y_pred[i][1])
                                    

                                    if current_close>previous_close and rsi<50 and sma_10>sma_30 and (pred_close-current_open)<(added_lag):
                                        if pred_close>=current_close and (current_open-current_low)<stop_loss_away:
                                            total_expected_wins=total_expected_wins+1
                                            stop_loss=current_open-current_low
                                            if stop_loss>biggest_buy_stop_loss:
                                                biggest_buy_stop_loss=stop_loss
                                        else:
                                            total_expected_losses=total_expected_losses+1
                                            
                                        total_expected_trade=total_expected_trade+1
                                    elif current_close<previous_close and rsi>50 and sma_10<sma_30  and (current_open-pred_close)<(added_lag):
                                        
                                        if pred_close<=current_close and (current_high-current_open)<stop_loss_away:
                                            total_expected_wins=total_expected_wins+1
                                            stop_loss=current_high-current_open
                                            if stop_loss>biggest_sell_stop_loss:
                                                biggest_sell_stop_loss=stop_loss
                                        else:
                                            total_expected_losses=total_expected_losses+1
                                        total_expected_trade=total_expected_trade+1
                                if (total_expected_wins/total_expected_losses)>best_profit_to_loss_ratio:# and total_expected_trade>overall_total_number_of_trades:
                                    best_profit_to_loss_ratio=(total_expected_wins/total_expected_losses)
                                    best_sma_10=sma_10_period
                                    best_sma_30=sma_30_period
                                    best_added_lag=added_lag
                                    best_rsi_period=rsi_period
                                    overall_total_number_of_trades=total_expected_trade
                                    print(f"Best Expected Profit to Loss Ratio: {best_profit_to_loss_ratio}")
                                    print(f"Total Expected Wins: {total_expected_wins}")
                                    print(f"Total Expected Losses: {total_expected_losses}")

                        except Exception as e:
                            #print(f'{symbol} failed')
                            raise e
                            pass
        doc = Document()
        doc.add_heading(f'Simulation Analysis for {symbol} in {timeframe} timeframe with Multiple of {added_lag}', level=1)

        # Add simulation results
        doc.add_heading('Best Results:', level=2)
        doc.add_paragraph(f'Best Expected Profit to Loss Ratio: {best_profit_to_loss_ratio}')
        doc.add_paragraph(f'Profit per trade: {added_lag*2}')
        doc.add_paragraph(f'Loss per trade : {added_lag}')
        doc.add_paragraph(f'best_sma_10:  {best_sma_10}')
        doc.add_paragraph(f'best_sma_30:   {best_sma_30}')
        doc.add_paragraph(f'Best rsi_period:   {rsi_period}')
        doc.add_paragraph(f'overall total_expected_trades:   {overall_total_number_of_trades}')
        doc.add_paragraph(f'Over A Span of {over_aspan_of} candles')
        # Save the document
        doc_path = f'{symbol}{added_lag}_simulation_analysis.docx'
        doc.save(doc_path)

        print(f'Document saved as {doc_path}')
asyncio.run(main2(timeframe,pages))