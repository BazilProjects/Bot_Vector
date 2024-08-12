import warnings
warnings.filterwarnings("ignore")
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split, learning_curve
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import PolynomialFeatures, OneHotEncoder
from sklearn.compose import ColumnTransformer
import plotly.graph_objs as go
import plotly.io as pio
from itertools import combinations
from metaapi_cloud_sdk import MetaApi
import asyncio
import joblib
import pandas_ta as ta
import sklearn
from sklearn.model_selection import cross_val_score
import os
from sklearn.decomposition import PCA
from sklearn.metrics import classification_report, brier_score_loss, accuracy_score, precision_score, recall_score, f1_score, roc_auc_score

from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier, ExtraTreesClassifier, GradientBoostingClassifier, HistGradientBoostingClassifier, AdaBoostClassifier, BaggingClassifier
from datetime import datetime, timedelta
token = os.getenv('TOKEN') or 'eyJhbGciOiJSUzUxMiIsInR5cCI6IkpXVCJ9.eyJfaWQiOiJlNWZkNGRkYWZmZmIyMDk2YTAyMWYzNjZiY2YxYjYwYSIsInBlcm1pc3Npb25zIjpbXSwiYWNjZXNzUnVsZXMiOlt7ImlkIjoidHJhZGluZy1hY2NvdW50LW1hbmFnZW1lbnQtYXBpIiwibWV0aG9kcyI6WyJ0cmFkaW5nLWFjY291bnQtbWFuYWdlbWVudC1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZXN0LWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1ycGMtYXBpIiwibWV0aG9kcyI6WyJtZXRhYXBpLWFwaTp3czpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZWFsLXRpbWUtc3RyZWFtaW5nLWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6d3M6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6Im1ldGFzdGF0cy1hcGkiLCJtZXRob2RzIjpbIm1ldGFzdGF0cy1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoicmlzay1tYW5hZ2VtZW50LWFwaSIsIm1ldGhvZHMiOlsicmlzay1tYW5hZ2VtZW50LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJjb3B5ZmFjdG9yeS1hcGkiLCJtZXRob2RzIjpbImNvcHlmYWN0b3J5LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJtdC1tYW5hZ2VyLWFwaSIsIm1ldGhvZHMiOlsibXQtbWFuYWdlci1hcGk6cmVzdDpkZWFsaW5nOio6KiIsIm10LW1hbmFnZXItYXBpOnJlc3Q6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6ImJpbGxpbmctYXBpIiwibWV0aG9kcyI6WyJiaWxsaW5nLWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19XSwidG9rZW5JZCI6IjIwMjEwMjEzIiwiaW1wZXJzb25hdGVkIjpmYWxzZSwicmVhbFVzZXJJZCI6ImU1ZmQ0ZGRhZmZmYjIwOTZhMDIxZjM2NmJjZjFiNjBhIiwiaWF0IjoxNzE4NjAwMTc5fQ.P_Q7S9gllY-A0ygrF72pmpyUrno0VZ6_gBXIT31fLOwKFcEEeqopfcEH7yUL6upwTe69YAitfuy3OcDfjBNL7D7Vnuh1FUUUb2EbuGEnBi-B3GSazhZ83uSeAa89zutuNsr7DrptYf-ZHIUw10NSbIHZjhTKOsT9GoV-lv6QLsaxe87wJ8hbv5ajWvo2VKYhJKv0fQFDP2SwsboGnjC4ioqMGoGFhAv2BNLs3nXZL8SWe0tEYIcLmfWB1sWTlPDrsnaOzRdyOzUwbRljoJzns1BUyHW375eQTR93oNjN_P1zJNL9J_V8rqc-nIJcQPhIFytwaFDy-Z3DJHZQJk6mAbDLkphjQAvGeXrwjjk_uQNuY0WzmJmv21dFeQ3aFJhW-wH7l_KYEYjwqxP_H2lU0_AXpiGxn0ZWOoRbyp2Uer8X2hBg_psUd7RlhmfVqXQqTkJLpFpZrP980rq1S_deFWWiZPihXqkzTsFwUTl2DWMyuTxwsQnCBi-Dbt11XY-1kXe7uABqL9L6YbaclohAUVQF1pIVONbELTkTzUqfzqf30TLsxT67xScvd51-smHTDBYyvSDjzWLQ-HCl9gMJlZ43NcJtA9fckE47yf7O5YaaaA3QXhNSO86jiDvJ_spVijZXi62P_d4d71fCpUZTum1WSwRAiF5scvIjwFMSL98'
accountId = os.getenv('ACCOUNT_ID') or '7416410e-1803-4778-bead-73b66d695bb5'
#symbol_list =['NZDUSDm','GBPTRYm','XAUUSDm','XAGUSDm',]


symbol_list = [
    'XAUUSDm',  # Gold/US Dollar (Commodity)
    'GBPAUDm' ,  # British Pound/Australian Dollar (Minor)
    'BTCUSDm',
    'EURUSDm',  # Euro/US Dollar (Major)
    'GBPUSDm',  # British Pound/US Dollar (Major)
    

]
timeframe='15m'
pages=5
n_estimators=7
min_samples_leaf=1
max_depth=50



def decimal_places(number):
    # Convert the number to a string
    num_str = str(number)
    
    # Check if there is a decimal point
    if '.' in num_str:
        # Find the index of the decimal point
        decimal_index = num_str.index('.')
        
        # Count the characters after the decimal point
        num_decimal_places = len(num_str) - decimal_index - 1
        
        return num_decimal_places
    else:
        # If there is no decimal point, return 0
        return 0
import json
from docx import Document

# Dictionary to store the metrics
performance_metrics = {}

# Function to add metrics to the dictionary
def add_metrics(symbol, model_name, mse, r2, mae, max_err):
    if symbol not in performance_metrics:
        performance_metrics[symbol] = {}
    performance_metrics[symbol][model_name] = {
        "Mean Squared Error": mse,
        "R2 Score": r2,
        "Mean Absolute Error": mae,
        "Maximum Error": max_err
    }


async def get_candles_m(timeframe,symbol,pages):
    api = MetaApi(token)
    account = await api.metatrader_account_api.get_account(accountId)
    initial_state = account.state
    deployed_states = ['DEPLOYING', 'DEPLOYED']
    timeframe=timeframe
    symbol=symbol
    if initial_state not in deployed_states:
        # wait until account is deployed and connected to broker
        print('Deploying account')
        await account.deploy()
        print('Waiting for API server to connect to broker (may take a few minutes)')
        await account.wait_connected()
        
    try:
        # Create an empty dataframe to store the candlestick data
        df = pd.DataFrame()


        # retrieve last 10K100
        pages = pages
        #print(f'Downloading {pages}K latest candles for {symbol}')
        started_at = datetime.now().timestamp()
        start_time = None
        candles = None
        for i in range(pages):
            try:
                # the API to retrieve historical market data is currently available for G1 only
                candles = await account.get_historical_candles(symbol=symbol, timeframe=timeframe, start_time=start_time)
                print(f'Downloaded {i}000 historical candles for {symbol}')
                
                
                if not candles:
                    pass
                else:
                    #Create a new dataframe for each iteration and add it to the main dataframe
                    new_df = pd.DataFrame(candles)
                    df = pd.concat([df, new_df], ignore_index=True)
                    print(f'Candles added to dataframe')
                df.to_csv(f'{symbol}{timeframe}{str(pages)}.csv', index=False)
            except:
                pass

        return df

    except Exception as e:
        raise e
        
async def main2(timeframe,pages):
    print('Up and runing')
    for symbol in symbol_list:
        try:

            df=pd.read_csv(f'../COLLECT CANDLES/{timeframe}/{symbol}{timeframe}{str(pages)}.csv')

            if not df.empty:
                df=df.drop(columns=['symbol','timeframe','brokerTime'])

                df['time'] = pd.to_datetime(df['time'])
                df['time'] = df['time'].astype(int) #// 10**99 
                df_new=df
                print(len(df))
                #df_new=prepare(df)
                #dt=pd.DataFrame(df_new)
                #print(dt)
                #print(dt)
                # Print all combinations
                df_new['next_close'] = df_new['close'].shift(-1)
                df_new['next_open']=df_new['open'].shift(-1)

                # Create labels: 1 for 'buy' (next close is higher), 0 for 'sell' (next close is lower)
                df_new['Label_close'] = (df_new['next_close'] > df_new['close']).astype(int)
                df_new['Label_open'] = (df_new['next_open'] > df_new['next_close']).astype(int)
 
                df_new=df_new.dropna()

                X =df_new.drop(columns=['next_close','next_open','Label_open','Label_close',])
                y= df_new[['Label_close','Label_open']]


                # Split the data into training and testing sets
                X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.01, random_state=42, shuffle=False)
                
                
                
                # Define a dictionary of models
                models = {
                    'DecisionTree': DecisionTreeClassifier(),
                    'ExtraTrees': ExtraTreesClassifier(n_estimators=1, random_state=42),
                    #'HistGradientBoosting': HistGradientBoostingClassifier(max_iter=500),
                    #'Bagging': BaggingClassifier(base_estimator=DecisionTreeClassifier(), n_estimators=10),
                }


                # Train each model, make predictions, and calculate metrics
                results = {}
                for name, model in models.items():
                    """
                    print(symbol)

                    scores = cross_val_score(model, X, y, cv=300, scoring='accuracy')  # You can use other scoring metrics as needed
                    print(f"{name} Cross-Validation Accuracy: {scores.mean()} ± {scores.std()}")
                    print(f' '*500)
                    """
                    print(symbol)

                    # Fit the model
                    model.fit(X_train, y_train)
                    
                    # Predict on the test set
                    predictions = model.predict(X_test)
                    # Calculate metrics
                    # Calculate additional evaluation metrics
                    accuracy_svm = accuracy_score(y_test, predictions)
                    precision_svm = precision_score(y_test, predictions, average='weighted')
                    recall_svm = recall_score(y_test, predictions, average='weighted')
                    f1_svm = f1_score(y_test, predictions, average='weighted')

                    # Print additional evaluation metrics
                    print("Additional Evaluation Metrics:")
                    print("Accuracy:", accuracy_svm)
                    print("Precision:", precision_svm)
                    print("Recall:", recall_svm)
                    print("F1 Score:", f1_svm)
                    model.fit(X, y)
                    joblib.dump((model, sklearn.__version__), f'Classifiers/{timeframe}/{name}{symbol}.pkl')
                    next_prediction=model.predict(np.array(X.iloc[-1]).reshape(1, -1))
                    if next_prediction[0][0]==1 and next_prediction[0][1]==0:
                        print('Buy')
                    if next_prediction[0][0]==0 and next_prediction[0][1]==1:
                        print('Sell')
                    add_metrics(symbol, model, precision_svm, accuracy_score, recall_svm, f1_svm)
                    
        except Exception as e:
            #print(f'{symbol} failed')
            #raise e
            pass
    doc = Document()
    doc.add_heading('Model Performance Metrics', 0)

    for symbol, models in performance_metrics.items():
        doc.add_heading(symbol, level=1)
        for model_name, metrics in models.items():
            doc.add_heading(f"Model: {model_name}", level=2)
            doc.add_paragraph(f"Mean Squared Error: {metrics['Mean Squared Error']}")
            #doc.add_paragraph(f"R2 Score: {metrics['R2 Score']:.2f}%")
            doc.add_paragraph(f"Mean Absolute Error: {metrics['Mean Absolute Error']}")
            doc.add_paragraph(f"Maximum Error: {metrics['Maximum Error']}")
            doc.add_paragraph("-" * 30)
            
    doc.save(f'model_performance_metrics{timeframe}.docx')

    # Optional: Print the contents for verification
    #print(json.dumps(performance_metrics, indent=4))
                
        
asyncio.run(main2(timeframe,pages))