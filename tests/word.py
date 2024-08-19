"""import docx

# Load the document
doc_path = "/home/omenyo/Documents/GitHub/Bot_Vector/tests/EURUSDm0.0008_simulation_analysis.docx"
doc = docx.Document(doc_path)

# Initialize an empty dictionary to store the extracted values
extracted_data = {}

# Loop through each paragraph and extract the relevant values
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    
    if "best_sma_10:" in text:
        extracted_data['best_sma_10'] = int(text.split(':')[-1].strip())
    
    elif "best_sma_30:" in text:
        extracted_data['best_sma_30'] = int(text.split(':')[-1].strip())
    
    elif "Best rsi_period:" in text:
        extracted_data['Best rsi_period'] = int(text.split(':')[-1].strip())

# Print the extracted data to verify
print("Extracted Data:", extracted_data)

# Example of how to access each value later
best_sma_10 = extracted_data.get('best_sma_10')
best_sma_30 = extracted_data.get('best_sma_30')
best_rsi_period = extracted_data.get('Best rsi_period')

print(f"best_sma_10: {best_sma_10}")
print(f"best_sma_30: {best_sma_30}")
print(f"Best rsi_period: {best_rsi_period}")
"""
import os
import docx

def extract_simulation_values(currency_pair_file):
    # Load the .docx file
    doc = docx.Document(currency_pair_file)
    
    # Initialize variables to store extracted values
    best_sma_10 = None
    best_sma_30 = None
    best_rsi_period = None
    
    # Iterate through each line in the document
    for para in doc.paragraphs:
        line = para.text.strip()
        
        # Extract values based on specific keywords
        if "best_sma_10" in line:
            best_sma_10 = line.split(":")[1].strip()
        elif "best_sma_30" in line:
            best_sma_30 = line.split(":")[1].strip()
        elif "Best rsi_period" in line:
            best_rsi_period = line.split(":")[1].strip()
    
    return best_sma_10, best_sma_30, best_rsi_period

# Directory where the .docx files are stored
directory = "/home/omenyo/Documents/GitHub/Bot_Vector/tests/"

currency_pairs = [
    "GBPUSDm0.001_simulation_analysis.docx",
    "EURUSDm0.0007_simulation_analysis.docx",
    "BTCUSDm600_simulation_analysis.docx",
    "XAUUSDm4_simulation_analysis.docx"
]


# Loop through each currency pair file
for pair_file in currency_pairs:
    full_path = os.path.join(directory, pair_file)
    
    # Extract the values from the file
    best_sma_10, best_sma_30, best_rsi_period = extract_simulation_values(full_path)
    
    # Print the results
    print(f"Results for {pair_file}:")
    print(f"best_sma_10: {best_sma_10}")
    print(f"best_sma_30: {best_sma_30}")
    print(f"Best rsi_period: {best_rsi_period}")
    print("-" * 30)
