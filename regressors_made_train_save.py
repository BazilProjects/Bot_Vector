import warnings
warnings.filterwarnings("ignore")
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split, learning_curve
from sklearn.linear_model import LinearRegression
from sklearn.tree import DecisionTreeRegressor
from sklearn.ensemble import RandomForestRegressor, ExtraTreesRegressor
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import PolynomialFeatures, OneHotEncoder
from sklearn.compose import ColumnTransformer
import plotly.graph_objs as go
import plotly.io as pio
from itertools import combinations
from metaapi_cloud_sdk import MetaApi
import asyncio
import joblib
import sklearn
import os
from docx import Document



from datetime import datetime, timedelta
token = os.getenv('TOKEN') or 'eyJhbGciOiJSUzUxMiIsInR5cCI6IkpXVCJ9.eyJfaWQiOiJlNWZkNGRkYWZmZmIyMDk2YTAyMWYzNjZiY2YxYjYwYSIsInBlcm1pc3Npb25zIjpbXSwiYWNjZXNzUnVsZXMiOlt7ImlkIjoidHJhZGluZy1hY2NvdW50LW1hbmFnZW1lbnQtYXBpIiwibWV0aG9kcyI6WyJ0cmFkaW5nLWFjY291bnQtbWFuYWdlbWVudC1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZXN0LWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1ycGMtYXBpIiwibWV0aG9kcyI6WyJtZXRhYXBpLWFwaTp3czpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZWFsLXRpbWUtc3RyZWFtaW5nLWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6d3M6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6Im1ldGFzdGF0cy1hcGkiLCJtZXRob2RzIjpbIm1ldGFzdGF0cy1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoicmlzay1tYW5hZ2VtZW50LWFwaSIsIm1ldGhvZHMiOlsicmlzay1tYW5hZ2VtZW50LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJjb3B5ZmFjdG9yeS1hcGkiLCJtZXRob2RzIjpbImNvcHlmYWN0b3J5LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJtdC1tYW5hZ2VyLWFwaSIsIm1ldGhvZHMiOlsibXQtbWFuYWdlci1hcGk6cmVzdDpkZWFsaW5nOio6KiIsIm10LW1hbmFnZXItYXBpOnJlc3Q6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6ImJpbGxpbmctYXBpIiwibWV0aG9kcyI6WyJiaWxsaW5nLWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19XSwidG9rZW5JZCI6IjIwMjEwMjEzIiwiaW1wZXJzb25hdGVkIjpmYWxzZSwicmVhbFVzZXJJZCI6ImU1ZmQ0ZGRhZmZmYjIwOTZhMDIxZjM2NmJjZjFiNjBhIiwiaWF0IjoxNzE4NjAwMTc5fQ.P_Q7S9gllY-A0ygrF72pmpyUrno0VZ6_gBXIT31fLOwKFcEEeqopfcEH7yUL6upwTe69YAitfuy3OcDfjBNL7D7Vnuh1FUUUb2EbuGEnBi-B3GSazhZ83uSeAa89zutuNsr7DrptYf-ZHIUw10NSbIHZjhTKOsT9GoV-lv6QLsaxe87wJ8hbv5ajWvo2VKYhJKv0fQFDP2SwsboGnjC4ioqMGoGFhAv2BNLs3nXZL8SWe0tEYIcLmfWB1sWTlPDrsnaOzRdyOzUwbRljoJzns1BUyHW375eQTR93oNjN_P1zJNL9J_V8rqc-nIJcQPhIFytwaFDy-Z3DJHZQJk6mAbDLkphjQAvGeXrwjjk_uQNuY0WzmJmv21dFeQ3aFJhW-wH7l_KYEYjwqxP_H2lU0_AXpiGxn0ZWOoRbyp2Uer8X2hBg_psUd7RlhmfVqXQqTkJLpFpZrP980rq1S_deFWWiZPihXqkzTsFwUTl2DWMyuTxwsQnCBi-Dbt11XY-1kXe7uABqL9L6YbaclohAUVQF1pIVONbELTkTzUqfzqf30TLsxT67xScvd51-smHTDBYyvSDjzWLQ-HCl9gMJlZ43NcJtA9fckE47yf7O5YaaaA3QXhNSO86jiDvJ_spVijZXi62P_d4d71fCpUZTum1WSwRAiF5scvIjwFMSL98'
accountId = os.getenv('ACCOUNT_ID') or '653d65c4-a70f-49ac-a6de-deea63238808'

symbol_list = [
    'XAUUSDm',  # Gold/US Dollar (Commodity)
    'GBPAUDm' ,  # British Pound/Australian Dollar (Minor)
    'BTCUSDm',
    #'EURUSDm',  # Euro/US Dollar (Major)
    #'GBPUSDm',  # British Pound/US Dollar (Major)
    'XAGUSDm',  # Silver/US Dollar (Commodity)
    #'AUDCHFm',  # Australian Dollar/Swiss Franc (Minor)
    #'NZDUSDm',  # New Zealand Dollar/US Dollar (Major)
    #'GBPTRYm',  # British Pound/Turkish Lira (Exotic)
    #'USDCHFm',  # US Dollar/Swiss Franc (Major)
    #'AUDUSDm',  # Australian Dollar/US Dollar (Major)
    #'EURGBPm',  # Euro/British Pound (Minor)
    #'GBPCHFm',  # British Pound/Swiss Franc (Minor)
    #'AUDJPYm',  # Australian Dollar/Japanese Yen (Minor)
    #'AUDNZDm',  # Australian Dollar/New Zealand Dollar (Minor)
    ##'EURCHFm',  # Euro/Swiss Franc (Minor)
    #'EURAUDm',  # Euro/Australian Dollar (Minor)
    #'EURCADm',  # Euro/Canadian Dollar (Minor)
    
    

]

data = {}
timeframe='15m'
pages=8
n_estimators=1
min_samples_leaf=1
shuffle=True
max_depth=50
test_size=0.01
def generate_new_features(df,column1,column2):
    # Check if required columns are present
    required_columns = [str(column1),str(column2)]#[f'{column1}', f'{column2}']
    for col in required_columns:
        if col not in df.columns:
            raise ValueError(f"DataFrame must contain '{col}' column.")

    # Addition/Subtraction
    df[f'{column1}_plus_{column2}'] = df[column1] + df[column2]
    df[f'{column1}_minus_{column2}'] = df[column1] - df[column2]
    #df[f'close_dif{column1}']=df[column1]//df[column1].min()
    #df[f'close_dif{column2}']=df[column2]//df[column2].min()
    # Multiplication/Division
    df[f'{column1}_times_{column2}'] = df[column1] * df[column2]
    df[f'{column1}_div_{column2}'] = df[column1] / df[column2]

    # Mean and Standard Deviation
    df[f'mean_column_{column1}{column2}'] = df[[column1, column2]].mean(axis=1)
    df[f'std_column_{column1}{column2}'] = df[[column1, column2]].std(axis=1)

    # Sum
    df[f'sum_column_{column1}{column2}'] = df[[column1, column2]].sum(axis=1)

    # Logarithm and Exponential
    

    df[column1] = df[column1].apply(lambda x: x if x > 1e-40 else 1e-40)  # Replace non-positive values
    df[column2] = df[column2].apply(lambda x: x if x > 1e-40 else 1e-40)  # Replace non-positive values

    # Trigonometric Functions
   
    df[f'log_column1_{column1}{column2}'] = np.log(df[column1])
    df[f'exp_column1_{column1}{column2}'] = np.exp(df[column1])

    # Square Root
    df[f'sqrt_column1_{column1}{column2}'] = np.sqrt(df[column1])
    #df[column2] = df[column2].apply(lambda x: x if x > 1e-30 else 1e-30)
    #df[f'sqrt_column1_2{column1}{column2}'] = np.sqrt(df[column2])

    # Rolling Mean and Cumulative Sum
    df[f'rolling_mean_column1_{column1}{column2}'] = df[column1].rolling(window=3).mean()
    df[f'cumulative_sum_column1_{column1}{column2}'] = df[column1].cumsum()

    #df[f'rolling_mean_column1_2{column1}{column2}'] = df[column2].rolling(window=3).mean()
    #df[f'cumulative_sum_column1_2{column1}{column2}'] = df[column2].cumsum()

    # Polynomial Features
    poly = PolynomialFeatures(degree=2)
    poly_features = poly.fit_transform(df[[column1, column2]])
    df[f'poly_column1_2_{column1}{column2}'] = poly_features[:, 2]  # x1 * x2 term
    df[f'poly_column1_2_squared_{column1}{column2}'] = poly_features[:, 3]  # x1^2 term
    df[f'poly_column2_squared_{column1}{column2}'] = poly_features[:, 4]  # x2^2 term

    # Interaction Term
    df[f'interaction_{column1}{column2}'] = df[column1] * df[column2]
    #if ('close' in [column1,column2]) and ('open' in [column1,column2]):
    # Binning
    df[f'binned_column_{column1}{column2}'] = pd.cut(df[column1], bins=3, labels=['low', 'medium', 'high'])
    #df[f'binned_column_2_{column1}{column2}'] = pd.cut(df[column2], bins=3, labels=['low', 'medium', 'high'])

    # Lagged Features
    df[f'lagged_column1_{column1}{column2}'] = df[column1].shift(1)
    df[f'lagged_column2_{column1}{column2}'] = df[column2].shift(1)
    #df[f'lagged_column3_{column1}{column2}'] = df[column1].shift(2)

    # Custom Function Example
    df[f'custom_column1_{column1}{column2}'] = df[column1].apply(lambda x: x * 2)
    #df[f'custom_column2_{column1}{column2}'] = df[column2].apply(lambda x: x * 2)
     # Adding trigonometric function results to DataFrame
    df[f'sin_close_{column1}{column2}'] =  np.sin(df[column1])
    df[f'sin_open_{column1}{column2}'] = np.sin(df[column2])

    df[f'cos_close_{column1}{column2}'] = np.cos(df[column1])
    df[f'cos_open_{column1}{column2}'] = np.cos(df[column2])

    df[f'tan_close_{column1}{column2}'] = np.tan(df[column1])
    df[f'tan_open_{column1}{column2}'] = np.tan(df[column2])



    # Adding hyperbolic function results to DataFrame
    df[f'sinh_close_{column1}{column2}'] = np.sinh(df[column1])
    df[f'sinh_open_{column1}{column2}'] = np.sinh(df[column2])

    df[f'cosh_close_{column1}{column2}'] = np.cosh(df[column1])
    df[f'cosh_open_{column1}{column2}'] = np.cosh(df[column2])

    df[f'tanh_close_{column1}{column2}'] = np.tanh(df[column1])
    df[f'tanh_open_{column1}{column2}'] = np.tanh(df[column2])

    return df
def decimal_places(number):
    # Convert the number to a string
    num_str = str(number)
    
    # Check if there is a decimal point
    if '.' in num_str:
        # Find the index of the decimal point
        decimal_index = num_str.index('.')
        
        # Count the characters after the decimal point
        num_decimal_places = len(num_str) - decimal_index - 1
        
        return num_decimal_places
    else:
        # If there is no decimal point, return 0
        return 0


# Function to add nested dictionary content to a .docx file
def add_nested_dict_to_docx(info_dict, filename=f'output{pages}{timeframe}.docx'):
    doc = Document()
    doc.add_heading('AI analysis and market forecast For timeframe', 0)
    for symbol, details in info_dict.items():
        doc.add_heading(symbol, level=1)
        for key, value in details.items():
            doc.add_paragraph(f"{key}: {value}", style='List Number')
    doc.save(filename)
    print(f"Information added to {filename} successfully.")


def add_stop_losse(df):
    pass
    # Initialize new columns
    df['type'] = None
    df['sell_high'] = None
    df['buy_low'] = None

    # Determine the candle type dynamically
    for i in range(df.index[1], len(df)):
        if df.at[i, 'close'] > df.at[i-1, 'close']:
            df.at[i, 'type'] = 'buy'
        else:
            df.at[i, 'type'] = 'sell'

    # Set the first candle type as 'buy' or 'sell' as needed
    df.at[0, 'type'] = 'buy' if df.at[0, 'close'] > df.at[0, 'open'] else 'sell'


    # Loop through the DataFrame to fill Sell_High and Buy_Low
    for index, row in df.iterrows():
        if row['type'] == 'sell':
            df.at[index, 'sell_high'] = row['high']
        elif row['type'] == 'buy':
            df.at[index, 'buy_low'] = row['low']
    # Loop through the DataFrame to fill Sell_High and Buy_Low
    for index, row in df.iterrows():
        if row['type'] == 'sell':
            df.at[index, 'sell_low'] = row['low']
        elif row['type'] == 'buy':
            df.at[index, 'buy_high'] = row['high']
    # Combine Sell_High and Buy_Low into one column
    df['stop_losses'] = df['sell_high'].combine_first(df['buy_low'])
    df['trade_max'] = df['sell_low'].combine_first(df['buy_high'])
    df['sell_high']= df['sell_high'].fillna(1)
    df['buy_low']= df['buy_low'].fillna(1)
    df['sell_low']= df['sell_low'].fillna(1)
    df['buy_high']= df['buy_high'].fillna(1)
    df['return'] = df['close'].pct_change()
    df['volatility'] = df['return'].rolling(window=10).std()
    df['moving_avg'] = df['close'].rolling(window=10).mean()
    # Drop the individual Sell_High and Buy_Low columns
    #df.drop(columns=['sell_high', 'buy_low',], inplace=True)
    return df
# Function to add data to the dictionary
def add_symbol_data(symbol, r2_2, mse_2, r2_2_1, mse_2_1, r2_2_2, mse_2_2, next_close, next_low, next_high, actual_close, actual_low, actual_high, diff2, diff1):
    if symbol not in data:
        data[symbol] = {}
    
    data[symbol] = {
        'r2_2': r2_2,
        'mse_2': mse_2,
        'r2_2_1': r2_2_1,
        'mse_2_1': mse_2_1,
        'r2_2_2': r2_2_2,
        'mse_2_2': mse_2_2,
        'next_close': next_close,
        'next_low': next_low,
        'next_high': next_high,
        'actual': {
            'close': actual_close,
            'low': actual_low,
            'high': actual_high
        },
        'diff2': diff2,
        'diff1': diff1
    }

def prepare(df):
    df=add_stop_losse(df)
    df=df.drop(columns=['symbol','timeframe','brokerTime'])

    df['time'] = pd.to_datetime(df['time'])
    df['time'] = df['time'].astype(int) #// 10**99 
    # Generate new features and handle data preparation
    df_new =df
    #print(df_new)
    #print(df.columns)
    # Columns to consider
    columns = ['high', 'low', 'close', 'open', 'tickVolume','time','stop_losses','trade_max','sell_high','sell_low','buy_low','buy_high']

    # Generate combinations of 2 pairs
    combinations_2 = list(combinations(columns, 2))

    # Print all combinations
    for comb in combinations_2:
        first_member, second_member = comb
        df_new = generate_new_features(df_new,first_member,second_member)
    
    time=df_new['time']
    df_new = df_new.dropna()  # Drop NaN values

    # Round numerical columns based on decimal places of the last 'close' value
    decimal_places = lambda number: len(str(number).split('.')[1])
    #df_new = df_new.round(decimal_places(df['close'].iloc[-1]))

    # Check for infinite values in each column
    inf_mask = df_new.isin([np.inf, -np.inf])

    # Drop columns with infinite values
    df_new = df_new.loc[:, ~inf_mask.any()]

    # Set a threshold for numerical columns
    threshold = 1e40  # Adjust the threshold as needed

    # Filter columns where all values are greater than the threshold
    columns_to_drop = df_new.select_dtypes(include=np.number).columns[(df_new.select_dtypes(include=np.number) > threshold).all()]
    #df_new['time']=time
    #print(columns_to_drop)
    #print(df_new.drop(columns=['Prediction']))
    df_new.drop(columns=columns_to_drop)
    return df_new
async def main2(timeframe,pages):
    print('Up and runing')
    for symbol in symbol_list:
        try:

            
            df=pd.read_csv(f'COLLECT CANDLES/{timeframe}/{symbol}{timeframe}{str(pages)}.csv')
            actual=df.tail(1)
            df=df.head(len(df)-1)
            

            #print(df)
            if not df.empty:
                df_new=prepare(df)
                dt=pd.DataFrame(df_new)
                #print(dt)
                #print(dt)
                # Print all combinations
                df_new['Candle_close'] = df_new['close'].shift(-1)
                df_new['Canlde_low']=df_new['low'].shift(-1)
                df_new['Canlde_high']=df_new['high'].shift(-1)

                #print(df_new[['Prediction','trade_max','stop_losses','stop_losses_predictions']])
                df_new=df_new.dropna()
                #print(df_new[['Prediction','trade_max','stop_losses','stop_losses_predictions']])
                
                X =df_new.drop(columns=['Candle_close','Canlde_low','Canlde_high'])
                y_close= df_new['Candle_close']
                y_low=df_new['Canlde_low']
                y_high=df_new['Canlde_high']
                print(len(X.columns))
                #print(X.columns)
                column_list=[]
                for column in X.columns:
                    try:
                        if str(column[:6])=='binned':
                            column_list.append(column)
                    except Exception as e:
                        pass
                for i in ['type']:
                    column_list.append(i)
                #print(column_list)
                #print(dt)
                transformer = ColumnTransformer(
                    transformers=[('cat', OneHotEncoder(), column_list)],
                    remainder='passthrough'
                )
                last_row= transformer.fit_transform(dt)
                last_row = np.array(last_row[-1]).reshape(1, -1)
                print(last_row)
                #print(last_row)
                # One-hot encode the 'binned_column1' feature
                transformer = ColumnTransformer(
                    transformers=[('cat', OneHotEncoder(), column_list)],
                    remainder='passthrough'
                )
                X = transformer.fit_transform(X)
                # Split the data into training and testing sets

                X_train, X_test, y_train, y_test = train_test_split(X, y_close, test_size=test_size, random_state=42, shuffle=shuffle)
                
                
                # Initialize RandomForestRegressor
                model_close =   DecisionTreeRegressor()#RandomForestRegressor(n_estimators=n_estimators,min_samples_leaf=min_samples_leaf,max_depth=max_depth,random_state=42)
                # Fit the RandomizedSearchCV
                model_close.fit(X_train, y_train)
                joblib.dump((model_close, sklearn.__version__), f'Regressors/Close/model{symbol}{timeframe}close.pkl')
                # Predict on the test set using the best model
                predictions_close = model_close.predict(X_test)
                
                y_pred=[round(pred, decimal_places(df['close'].iloc[-1])) for pred in predictions_close]
                # Calculate metrics
                mse_2 = mean_squared_error(y_test, y_pred)
                r2_2 = r2_score(y_test, y_pred) * 100
                print(f'{symbol}')
                
                diff1=y_test-y_pred
                #print(diff1)
                diff1=diff1.max()

                _, _, y_train, y_test = train_test_split(X, y_low, test_size=test_size, random_state=42, shuffle=shuffle)
                
                model_1 =    DecisionTreeRegressor()#RandomForestRegressor(n_estimators=n_estimators,min_samples_leaf=min_samples_leaf,max_depth=max_depth,random_state=42)
                # Fit the RandomizedSearchCV
                model_1.fit(X_train, y_train)
                joblib.dump((model_1, sklearn.__version__), f'Regressors/Low/model{symbol}{timeframe}low.pkl')
                # Predict on the test set using the best model
                predictions_1 = model_1.predict(X_test)
                y_pred=[round(pred, decimal_places(df['close'].iloc[-1])) for pred in predictions_1]
                # Calculate metrics
                mse_2_1 = mean_squared_error(y_test, y_pred)
                r2_2_1 = r2_score(y_test, y_pred) * 100

                diff2=y_test-y_pred
                diff2=diff2.min()
                print(f'Max difference between pred and actual for second model:{diff2}')
                
                


                # Get the last row of features
                # One-hot encode the 'binned_column1' feature
                #last_row = np.array(dt.iloc[-1]).reshape(1, -1)
                
                # Predict the next price
                next_close = model_close.predict(last_row)
                next_low = model_1.predict(last_row)

                _, _, y_train, y_test = train_test_split(X, y_high, test_size=test_size, random_state=42, shuffle=shuffle)
                
                model_1 =    DecisionTreeRegressor()#RandomForestRegressor(n_estimators=n_estimators,min_samples_leaf=min_samples_leaf,max_depth=max_depth,random_state=42)
                # Fit the RandomizedSearchCV
                model_1.fit(X_train, y_train)
                joblib.dump((model_1, sklearn.__version__), f'Regressors/High/model{symbol}{timeframe}high.pkl')
                # Predict on the test set using the best model
                predictions_1 = model_1.predict(X_test)
                y_pred=[round(pred, decimal_places(df['close'].iloc[-1])) for pred in predictions_1]
                # Calculate metrics
                mse_2_2 = mean_squared_error(y_test, y_pred)
                r2_2_2 = r2_score(y_test, y_pred) * 100
                next_high = model_1.predict(last_row)

                print(f"Testing R-squared Close: {r2_2}")
                print(f"Testing MSE Close: {mse_2}")
                print('-'*20)
                print(f"Testing R-squared Low: {r2_2_1}")
                print(f"Testing MSE Low: {mse_2_1}")
                print('-'*20)
                print(f"Testing R-squared High: {r2_2_2}")
                print(f"Testing MSE High: {mse_2_2}")
                print('-'*20)

                actual_close=actual['close'].values
                actual_low=actual['low'].values
                actual_high=actual['high'].values

                next_close=round(next_close[0],decimal_places(df['close'].iloc[-3]))
                next_low=round(next_low[0],decimal_places(df['close'].iloc[-3]))
                next_high=round(next_high[0],decimal_places(df['close'].iloc[-3]))
                #print(df_new['high'].iloc[-1])
                print(f"Next predicted close price: {next_close} and Actual : {actual_close}")
                print(f"Next predicted low price: {next_low} and Actual : {actual_low}")
                print(f"Next predicted high price: {next_high} and Actual : {actual_high}")


                diff2=y_test-y_pred
                diff2=diff2.min()
                print(f'Max difference between pred and actual for Third model:{diff2}')
                
                print(f'Max difference between pred and actual for first model:{diff1}')
                
                print(f'Number of Wining Trades')
                print('*'*20)
                print('*'*20)
                print('*'*20)

                add_symbol_data(
                    symbol, 
                    r2_2, mse_2, 
                    r2_2_1, mse_2_1, 
                    r2_2_2, mse_2_2, 
                    next_close, next_low, next_high, 
                    actual_close, actual_low, actual_high, 
                    diff2, diff1
                )
            

        except Exception as e:
            #print(f'{symbol} failed')
            raise e
            pass
        # Call the function to add the nested dictionary content to the .docx file
    add_nested_dict_to_docx(data)
asyncio.run(main2(timeframe,pages))