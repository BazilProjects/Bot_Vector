import warnings
warnings.filterwarnings("ignore")
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split, learning_curve
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor#,ExtraTressRegressor
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import PolynomialFeatures, OneHotEncoder
from sklearn.compose import ColumnTransformer
import plotly.graph_objs as go
import plotly.io as pio
from itertools import combinations
from metaapi_cloud_sdk import MetaApi
import asyncio
import joblib
import pandas_ta as ta
import sklearn
import os
from sklearn.decomposition import PCA
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error, max_error

from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier, ExtraTreesClassifier, GradientBoostingClassifier, HistGradientBoostingClassifier, AdaBoostClassifier, BaggingClassifier
from sklearn.model_selection import cross_val_score
from datetime import datetime, timedelta
token = os.getenv('TOKEN') or 'eyJhbGciOiJSUzUxMiIsInR5cCI6IkpXVCJ9.eyJfaWQiOiJlNWZkNGRkYWZmZmIyMDk2YTAyMWYzNjZiY2YxYjYwYSIsInBlcm1pc3Npb25zIjpbXSwiYWNjZXNzUnVsZXMiOlt7ImlkIjoidHJhZGluZy1hY2NvdW50LW1hbmFnZW1lbnQtYXBpIiwibWV0aG9kcyI6WyJ0cmFkaW5nLWFjY291bnQtbWFuYWdlbWVudC1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZXN0LWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1ycGMtYXBpIiwibWV0aG9kcyI6WyJtZXRhYXBpLWFwaTp3czpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZWFsLXRpbWUtc3RyZWFtaW5nLWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6d3M6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6Im1ldGFzdGF0cy1hcGkiLCJtZXRob2RzIjpbIm1ldGFzdGF0cy1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoicmlzay1tYW5hZ2VtZW50LWFwaSIsIm1ldGhvZHMiOlsicmlzay1tYW5hZ2VtZW50LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJjb3B5ZmFjdG9yeS1hcGkiLCJtZXRob2RzIjpbImNvcHlmYWN0b3J5LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJtdC1tYW5hZ2VyLWFwaSIsIm1ldGhvZHMiOlsibXQtbWFuYWdlci1hcGk6cmVzdDpkZWFsaW5nOio6KiIsIm10LW1hbmFnZXItYXBpOnJlc3Q6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6ImJpbGxpbmctYXBpIiwibWV0aG9kcyI6WyJiaWxsaW5nLWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19XSwidG9rZW5JZCI6IjIwMjEwMjEzIiwiaW1wZXJzb25hdGVkIjpmYWxzZSwicmVhbFVzZXJJZCI6ImU1ZmQ0ZGRhZmZmYjIwOTZhMDIxZjM2NmJjZjFiNjBhIiwiaWF0IjoxNzE4NjAwMTc5fQ.P_Q7S9gllY-A0ygrF72pmpyUrno0VZ6_gBXIT31fLOwKFcEEeqopfcEH7yUL6upwTe69YAitfuy3OcDfjBNL7D7Vnuh1FUUUb2EbuGEnBi-B3GSazhZ83uSeAa89zutuNsr7DrptYf-ZHIUw10NSbIHZjhTKOsT9GoV-lv6QLsaxe87wJ8hbv5ajWvo2VKYhJKv0fQFDP2SwsboGnjC4ioqMGoGFhAv2BNLs3nXZL8SWe0tEYIcLmfWB1sWTlPDrsnaOzRdyOzUwbRljoJzns1BUyHW375eQTR93oNjN_P1zJNL9J_V8rqc-nIJcQPhIFytwaFDy-Z3DJHZQJk6mAbDLkphjQAvGeXrwjjk_uQNuY0WzmJmv21dFeQ3aFJhW-wH7l_KYEYjwqxP_H2lU0_AXpiGxn0ZWOoRbyp2Uer8X2hBg_psUd7RlhmfVqXQqTkJLpFpZrP980rq1S_deFWWiZPihXqkzTsFwUTl2DWMyuTxwsQnCBi-Dbt11XY-1kXe7uABqL9L6YbaclohAUVQF1pIVONbELTkTzUqfzqf30TLsxT67xScvd51-smHTDBYyvSDjzWLQ-HCl9gMJlZ43NcJtA9fckE47yf7O5YaaaA3QXhNSO86jiDvJ_spVijZXi62P_d4d71fCpUZTum1WSwRAiF5scvIjwFMSL98'
accountId = os.getenv('ACCOUNT_ID') or '7416410e-1803-4778-bead-73b66d695bb5'
#symbol_list =['NZDUSDm','GBPTRYm','XAUUSDm','XAGUSDm',]


symbol_list = [
    'EURUSDm', 'GBPUSDm',
    
    

]
timeframe='15m'
pages=8
n_estimators=7
min_samples_leaf=1
max_depth=50

def generate_new_features(df,column1,column2):
    # Check if required columns are present
    required_columns = [str(column1),str(column2)]#[f'{column1}', f'{column2}']
    for col in required_columns:
        if col not in df.columns:
            raise ValueError(f"DataFrame must contain '{col}' column.")

    # Addition/Subtraction
    df[f'{column1}_plus_{column2}'] = df[column1] + df[column2]
    df[f'{column1}_minus_{column2}'] = df[column1] - df[column2]
    #df[f'close_dif{column1}']=df[column1]//df[column1].min()
    #df[f'close_dif{column2}']=df[column2]//df[column2].min()
    # Multiplication/Division
    df[f'{column1}_times_{column2}'] = df[column1] * df[column2]
    df[f'{column1}_div_{column2}'] = df[column1] / df[column2]

    # Mean and Standard Deviation
    df[f'mean_column_{column1}{column2}'] = df[[column1, column2]].mean(axis=1)
    df[f'std_column_{column1}{column2}'] = df[[column1, column2]].std(axis=1)

    # Sum
    df[f'sum_column_{column1}{column2}'] = df[[column1, column2]].sum(axis=1)

    # Logarithm and Exponential
    
    df[column1] = df[column1].apply(lambda x: x if x > 1e-40 else 1e-40)  # Replace non-positive values
    
    df[column2] = df[column2].apply(lambda x: x if x > 1e-40 else 1e-40)  # Replace non-positive values

    # Trigonometric Functions
   
    df[f'log_column1_{column1}{column2}'] = np.log(df[column1])
    df[f'exp_column1_{column1}{column2}'] = np.exp(df[column1])

    # Square Root
    df[f'sqrt_column1_{column1}{column2}'] = np.sqrt(df[column1])
    #df[column2] = df[column2].apply(lambda x: x if x > 1e-30 else 1e-30)
    #df[f'sqrt_column1_2{column1}{column2}'] = np.sqrt(df[column2])

    # Rolling Mean and Cumulative Sum
    df[f'rolling_mean_column1_{column1}{column2}'] = df[column1].rolling(window=3).mean()
    df[f'cumulative_sum_column1_{column1}{column2}'] = df[column1].cumsum()

    #df[f'rolling_mean_column1_2{column1}{column2}'] = df[column2].rolling(window=3).mean()
    #df[f'cumulative_sum_column1_2{column1}{column2}'] = df[column2].cumsum()

    # Polynomial Features
    poly = PolynomialFeatures(degree=2)
    poly_features = poly.fit_transform(df[[column1, column2]])
    df[f'poly_column1_2_{column1}{column2}'] = poly_features[:, 2]  # x1 * x2 term
    df[f'poly_column1_2_squared_{column1}{column2}'] = poly_features[:, 3]  # x1^2 term
    df[f'poly_column2_squared_{column1}{column2}'] = poly_features[:, 4]  # x2^2 term

    # Interaction Term
    df[f'interaction_{column1}{column2}'] = df[column1] * df[column2]
    #if ('close' in [column1,column2]) and ('open' in [column1,column2]):
    # Binning
    df[f'binned_column_{column1}{column2}'] = pd.cut(df[column1], bins=3, labels=['low', 'medium', 'high'])
    #df[f'binned_column_2_{column1}{column2}'] = pd.cut(df[column2], bins=3, labels=['low', 'medium', 'high'])

    # Lagged Features
    df[f'lagged_column1_{column1}{column2}'] = df[column1].shift(1)
    df[f'lagged_column2_{column1}{column2}'] = df[column2].shift(1)
    #df[f'lagged_column3_{column1}{column2}'] = df[column1].shift(2)

    # Custom Function Example
    df[f'custom_column1_{column1}{column2}'] = df[column1].apply(lambda x: x * 2)
    #df[f'custom_column2_{column1}{column2}'] = df[column2].apply(lambda x: x * 2)
     # Adding trigonometric function results to DataFrame
    df[f'sin_close_{column1}{column2}'] =  np.sin(df[column1])
    df[f'sin_open_{column1}{column2}'] = np.sin(df[column2])

    df[f'cos_close_{column1}{column2}'] = np.cos(df[column1])
    df[f'cos_open_{column1}{column2}'] = np.cos(df[column2])

    df[f'tan_close_{column1}{column2}'] = np.tan(df[column1])
    df[f'tan_open_{column1}{column2}'] = np.tan(df[column2])



    # Adding hyperbolic function results to DataFrame
    df[f'sinh_close_{column1}{column2}'] = np.sinh(df[column1])
    df[f'sinh_open_{column1}{column2}'] = np.sinh(df[column2])

    df[f'cosh_close_{column1}{column2}'] = np.cosh(df[column1])
    df[f'cosh_open_{column1}{column2}'] = np.cosh(df[column2])

    df[f'tanh_close_{column1}{column2}'] = np.tanh(df[column1])
    df[f'tanh_open_{column1}{column2}'] = np.tanh(df[column2])

    return df

def decimal_places(number):
    # Convert the number to a string
    num_str = str(number)
    
    # Check if there is a decimal point
    if '.' in num_str:
        # Find the index of the decimal point
        decimal_index = num_str.index('.')
        
        # Count the characters after the decimal point
        num_decimal_places = len(num_str) - decimal_index - 1
        
        return num_decimal_places
    else:
        # If there is no decimal point, return 0
        return 0
import json
from docx import Document

# Dictionary to store the metrics
performance_metrics = {}

# Function to add metrics to the dictionary
def add_metrics(symbol, model_name, mse, r2, mae, max_err):
    if symbol not in performance_metrics:
        performance_metrics[symbol] = {}
    performance_metrics[symbol][model_name] = {
        "Mean Squared Error": mse,
        "R2 Score": r2,
        "Mean Absolute Error": mae,
        "Maximum Error": max_err
    }

def add_stop_losse(df):
    pass
    # Initialize new columns
    df['type'] = None
    df['sell_high'] = None
    df['buy_low'] = None

    # Determine the candle type dynamically
    for i in range(df.index[1], len(df)):
        if df.at[i, 'close'] > df.at[i-1, 'close']:
            df.at[i, 'type'] = 'buy'
        else:
            df.at[i, 'type'] = 'sell'

    # Set the first candle type as 'buy' or 'sell' as needed
    df.at[0, 'type'] = 'buy' if df.at[0, 'close'] > df.at[0, 'open'] else 'sell'


    # Loop through the DataFrame to fill Sell_High and Buy_Low
    for index, row in df.iterrows():
        if row['type'] == 'sell':
            df.at[index, 'sell_high'] = row['high']
        elif row['type'] == 'buy':
            df.at[index, 'buy_low'] = row['low']
    # Loop through the DataFrame to fill Sell_High and Buy_Low
    for index, row in df.iterrows():
        if row['type'] == 'sell':
            df.at[index, 'sell_low'] = row['low']
        elif row['type'] == 'buy':
            df.at[index, 'buy_high'] = row['high']
    # Combine Sell_High and Buy_Low into one column
    df['stop_losses'] = df['sell_high'].combine_first(df['buy_low'])
    df['trade_max'] = df['sell_low'].combine_first(df['buy_high'])
    df['sell_high']= df['sell_high'].fillna(1)
    df['buy_low']= df['buy_low'].fillna(1)
    df['sell_low']= df['sell_low'].fillna(1)
    df['buy_high']= df['buy_high'].fillna(1)
    df['return'] = df['close'].pct_change()
    df['volatility'] = df['return'].rolling(window=10).std()
    df['moving_avg'] = df['close'].rolling(window=10).mean()
    # Drop the individual Sell_High and Buy_Low columns
    #df.drop(columns=['sell_high', 'buy_low',], inplace=True)
    return df

def prepare(df):
    df=add_stop_losse(df)
    df=df.drop(columns=['symbol','timeframe','brokerTime'])

    df['time'] = pd.to_datetime(df['time'])
    df['time'] = df['time'].astype(int) #// 10**99 
    # Generate new features and handle data preparation
    df_new =df
    #print(df_new)
    #print(df.columns)
    # Columns to consider
    columns = ['high', 'low', 'close', 'open', 'tickVolume','time','stop_losses','trade_max','sell_high','sell_low','buy_low','buy_high']

    # Generate combinations of 2 pairs
    combinations_2 = list(combinations(columns, 2))

    # Print all combinations
    for comb in combinations_2:
        first_member, second_member = comb
        df_new = generate_new_features(df_new,first_member,second_member)
    
    time=df_new['time']
    df_new = df_new.dropna()  # Drop NaN values

    # Round numerical columns based on decimal places of the last 'close' value
    decimal_places = lambda number: len(str(number).split('.')[1])
    #df_new = df_new.round(decimal_places(df['close'].iloc[-1]))

    # Check for infinite values in each column
    inf_mask = df_new.isin([np.inf, -np.inf])

    # Drop columns with infinite values
    df_new = df_new.loc[:, ~inf_mask.any()]

    # Set a threshold for numerical columns
    threshold = 1e30  # Adjust the threshold as needed

    # Filter columns where all values are greater than the threshold
    columns_to_drop = df_new.select_dtypes(include=np.number).columns[(df_new.select_dtypes(include=np.number) > threshold).all()]
    #df_new['time']=time
    #print(columns_to_drop)
    #print(df_new.drop(columns=['Prediction']))
    #df_new.drop(columns=columns_to_drop)
    return df_new

async def get_candles_m(timeframe,symbol,pages):
    api = MetaApi(token)
    account = await api.metatrader_account_api.get_account(accountId)
    initial_state = account.state
    deployed_states = ['DEPLOYING', 'DEPLOYED']
    timeframe=timeframe
    symbol=symbol
    if initial_state not in deployed_states:
        # wait until account is deployed and connected to broker
        print('Deploying account')
        await account.deploy()
        print('Waiting for API server to connect to broker (may take a few minutes)')
        await account.wait_connected()
        
    try:
        # Create an empty dataframe to store the candlestick data
        df = pd.DataFrame()


        # retrieve last 10K100
        pages = pages
        #print(f'Downloading {pages}K latest candles for {symbol}')
        started_at = datetime.now().timestamp()
        start_time = None
        candles = None
        for i in range(pages):
            try:
                # the API to retrieve historical market data is currently available for G1 only
                candles = await account.get_historical_candles(symbol=symbol, timeframe=timeframe, start_time=start_time)
                print(f'Downloaded {i}000 historical candles for {symbol}')
                
                
                if not candles:
                    pass
                else:
                    #Create a new dataframe for each iteration and add it to the main dataframe
                    new_df = pd.DataFrame(candles)
                    df = pd.concat([df, new_df], ignore_index=True)
                    print(f'Candles added to dataframe')
                df.to_csv(f'{symbol}{timeframe}{str(pages)}.csv', index=False)
            except:
                pass

        return df

    except Exception as e:
        raise e
        
async def main2(timeframe,pages):
    print('Up and runing')
    for symbol in symbol_list:
        try:

            df=pd.read_csv(f'COLLECT CANDLES/{timeframe}/{symbol}{timeframe}{str(pages)}.csv')

            if not df.empty:
                df_new=prepare(df)
                dt=pd.DataFrame(df_new)
                #print(dt)
                #print(dt)
                # Print all combinations
                df_new['Candle_close'] = df_new['close'].shift(-1)
                df_new['Canlde_low']=df_new['low'].shift(-1)
                df_new['Canlde_high']=df_new['high'].shift(-1)

                # Create labels: 1 for 'buy' (next close is higher), 0 for 'sell' (next close is lower)
                df_new['Label'] = (df_new['Candle_close'] > df_new['close']).astype(int)


                df_new=df_new.dropna()

                X =df_new.drop(columns=['Candle_close','Canlde_low','Canlde_high','Label'])
                y_close= df_new['Label']
                y_low=df_new['Canlde_low']
                y_high=df_new['Canlde_high']



                #print(X.columns)
                column_list=[]
                for column in X.columns:
                    try:
                        if str(column[:6])=='binned':
                            column_list.append(column)
                    except Exception as e:
                        pass
                for i in ['type']:
                    column_list.append(i)
                #print(column_list)
                #print(dt)
                transformer = ColumnTransformer(
                    transformers=[('cat', OneHotEncoder(), column_list)],
                    remainder='passthrough'
                )
                last_row= transformer.fit_transform(dt)
                last_row = np.array(last_row[-3:])#.reshape(1, -1)
                #print(last_row)
                from sklearn.preprocessing import StandardScaler

                scaler = StandardScaler()
                last_row = scaler.fit_transform(last_row)

                # Split the data into training and testing sets
                ""
                # One-hot encode the 'binned_column1' feature
                transformer = ColumnTransformer(
                    transformers=[('cat', OneHotEncoder(), column_list)],
                    remainder='passthrough'
                )
                X=transformer.fit_transform(X)

                """
                scaler = StandardScaler()
                X = scaler.fit_transform(X)
                """
                

                # Split the data into training and testing sets




                X_train, X_test, y_train, y_test = train_test_split(X, y_close, test_size=0.01, random_state=42, shuffle=False)
                
                
                
                # Define a dictionary of models
                models = {
                    'DecisionTree': DecisionTreeClassifier(),
                    'ExtraTrees': ExtraTreesClassifier(n_estimators=1, random_state=42),
                    #'HistGradientBoosting': HistGradientBoostingClassifier(max_iter=500),
                    #'Bagging': BaggingClassifier(base_estimator=DecisionTreeClassifier(), n_estimators=10),
                }


                # Train each model, make predictions, and calculate metrics
                results = {}
                for name, model in models.items():
                    print(symbol)

                    # Fit the model
                    model.fit(X_train, y_train)
                    joblib.dump((model, sklearn.__version__), f'Classifiers/{timeframe}/{name}{symbol}.pkl')
                    # Predict on the test set
                    y_pred = model.predict(X_test)
                    # Calculate metrics
                    mse = mean_squared_error(y_test, y_pred)
                    r2 = r2_score(y_test, y_pred) * 100
                    mae = mean_absolute_error(y_test, y_pred)
                    max_err = max_error(y_test, y_pred)
                    print(symbol)
                    print(f"Model: {name}")
                    print(f"Mean Squared Error: {mse}")
                    print(f"R2 Score: {r2:.2f}%")
                    print(f"Mean Absolute Error: {mae}")
                    print(f"Maximum Error: {max_err}")
                    print("-" * 30)
                    add_metrics(symbol, model, mse, r2, mae, max_err)
        except Exception as e:
            #print(f'{symbol} failed')
            #raise e
            pass
    doc = Document()
    doc.add_heading('Model Performance Metrics', 0)

    for symbol, models in performance_metrics.items():
        doc.add_heading(symbol, level=1)
        for model_name, metrics in models.items():
            doc.add_heading(f"Model: {model_name}", level=2)
            doc.add_paragraph(f"Mean Squared Error: {metrics['Mean Squared Error']}")
            doc.add_paragraph(f"R2 Score: {metrics['R2 Score']:.2f}%")
            doc.add_paragraph(f"Mean Absolute Error: {metrics['Mean Absolute Error']}")
            doc.add_paragraph(f"Maximum Error: {metrics['Maximum Error']}")
            doc.add_paragraph("-" * 30)
            
    doc.save(f'model_performance_metrics{timeframe}.docx')

    # Optional: Print the contents for verification
    #print(json.dumps(performance_metrics, indent=4))
                
        
asyncio.run(main2(timeframe,pages))